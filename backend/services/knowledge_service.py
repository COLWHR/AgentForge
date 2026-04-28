import hashlib
import re
import uuid
from io import BytesIO
from datetime import timezone
from typing import List
from zipfile import BadZipFile

from sqlalchemy import delete, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.models.orm import KnowledgeChunk, KnowledgeDocument
from backend.models.schemas import KnowledgeDocumentRead, KnowledgeSearchResult
from backend.services.knowledge_parser import knowledge_parser
from backend.services.knowledge_retrieval_ranker import knowledge_retrieval_ranker


class KnowledgeService:
    CHUNK_SIZE = 900
    CHUNK_OVERLAP = 120
    MAX_UPLOAD_BYTES = 100 * 1024 * 1024
    SUPPORTED_UPLOAD_EXTENSIONS = {".pdf", ".docx"}
    SUPPORTED_UPLOAD_MIME_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    SEARCH_STOPWORDS = {
        "a",
        "an",
        "and",
        "are",
        "do",
        "for",
        "how",
        "i",
        "is",
        "of",
        "or",
        "the",
        "to",
        "what",
        "你",
        "你们",
        "我",
        "我们",
        "他",
        "她",
        "它",
        "有",
        "是",
        "的",
        "了",
        "吗",
        "呢",
        "啊",
        "吧",
        "么",
        "什么",
        "怎么",
        "如何",
        "一下",
        "一个",
    }

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        tokens: list[str] = []
        for token in re.findall(r"[A-Za-z0-9_]+|[\u4e00-\u9fff]+", text.lower()):
            tokens.append(token)
            if re.search(r"[\u4e00-\u9fff]", token):
                tokens.extend(list(token))
                tokens.extend(token[index : index + 2] for index in range(0, max(len(token) - 1, 0)))
        return tokens

    @staticmethod
    def _contains_cjk(token: str) -> bool:
        return re.search(r"[\u4e00-\u9fff]", token) is not None

    @classmethod
    def _is_informative_token(cls, token: str) -> bool:
        normalized = token.strip().lower()
        if not normalized or normalized in cls.SEARCH_STOPWORDS:
            return False
        if cls._contains_cjk(normalized) and len(normalized) <= 1:
            return False
        if not cls._contains_cjk(normalized) and len(normalized) <= 1:
            return False
        return True

    @classmethod
    def _filter_query_tokens(cls, tokens: list[str]) -> list[str]:
        return [token for token in tokens if cls._is_informative_token(token)]

    @staticmethod
    def _token_weight(token: str) -> float:
        length = len(token.strip())
        if KnowledgeService._contains_cjk(token) and length >= 2:
            return 2.0
        if length >= 6:
            return 3.0
        if length >= 3:
            return 2.0
        return 1.0

    @classmethod
    def _split_content(cls, content: str) -> list[str]:
        return [chunk.content for chunk in knowledge_parser.parse(title="", content=content)]

    @classmethod
    def _split_articles(cls, content: str) -> list[str]:
        pattern = re.compile(r"(?=(第\s*[一二三四五六七八九十百0-9]+\s*条))")
        parts = pattern.split(content)
        if len(parts) < 3:
            return []
        chunks: list[str] = []
        prefix = parts[0].strip()
        if prefix and len(prefix) > 20:
            chunks.extend(cls._fallback_split(prefix))
        for index in range(1, len(parts), 2):
            label = parts[index].strip()
            body = parts[index + 1] if index + 1 < len(parts) else ""
            chunk = f"{label}{body}".strip()
            if chunk:
                chunks.extend(cls._fallback_split(chunk, preserve_article=True))
        return chunks

    @classmethod
    def _fallback_split(cls, content: str, *, preserve_article: bool = False) -> list[str]:
        if len(content) <= 1200:
            return [content.strip()]
        chunks: list[str] = []
        article_label = cls.extract_article_label(content) if preserve_article else None
        start = 0
        while start < len(content):
            end = min(start + cls.CHUNK_SIZE, len(content))
            chunk = content[start:end].strip()
            if chunk:
                if article_label and article_label not in chunk[:30]:
                    chunk = f"{article_label}\n{chunk}"
                chunks.append(chunk)
            if end >= len(content):
                break
            start = max(end - cls.CHUNK_OVERLAP, start + 1)
        return chunks

    @staticmethod
    def chinese_number_to_int(value: str) -> int | None:
        return knowledge_parser.chinese_number_to_int(value)

    @classmethod
    def extract_article_no(cls, text: str) -> str | None:
        return knowledge_parser.extract_article_no(text)

    @classmethod
    def extract_article_label(cls, text: str) -> str | None:
        return knowledge_parser.extract_article_label(text)

    @staticmethod
    def _as_iso(value: object) -> str:
        if hasattr(value, "astimezone"):
            return value.astimezone(timezone.utc).isoformat()  # type: ignore[union-attr]
        return ""

    async def create_document(
        self,
        db: AsyncSession,
        *,
        agent_id: uuid.UUID,
        team_id: uuid.UUID,
        title: str,
        content: str,
        source_filename: str | None = None,
        source_mime_type: str | None = None,
    ) -> KnowledgeDocumentRead:
        document = KnowledgeDocument(
            id=uuid.uuid4(),
            agent_id=agent_id,
            team_id=team_id,
            title=title.strip(),
            content=content.strip(),
            document_type=self._infer_document_type(title=title, content=content),
            source_filename=source_filename,
            source_mime_type=source_mime_type,
            source_hash=hashlib.sha256(content.strip().encode("utf-8")).hexdigest(),
            status="ACTIVE",
        )
        db.add(document)
        await db.flush()

        parsed_chunks = knowledge_parser.parse(title=document.title, content=document.content)
        article_count = len({chunk.article_no for chunk in parsed_chunks if chunk.article_no})
        document.document_metadata = {
            **(document.document_metadata or {}),
            "article_count": article_count,
            "chunk_count": len(parsed_chunks),
            "parser_version": "knowledge-parser-v1",
        }
        for index, chunk in enumerate(parsed_chunks):
            token_text = " ".join(self._tokenize(f"{document.title}\n{chunk.content}"))
            db.add(
                KnowledgeChunk(
                    id=uuid.uuid4(),
                    document_id=document.id,
                    agent_id=agent_id,
                    team_id=team_id,
                    title=document.title,
                    content=chunk.content,
                    chunk_index=index,
                    token_text=token_text,
                    chunk_type=chunk.chunk_type,
                    section_path=chunk.section_path,
                    article_no=chunk.article_no,
                    article_label=chunk.article_label,
                    page_no=chunk.page_no,
                    start_char=chunk.start_char,
                    end_char=chunk.end_char,
                    chunk_metadata=chunk.metadata,
                )
            )
        await db.commit()
        await db.refresh(document)
        return self._read_model(document, chunk_count=len(parsed_chunks))

    @staticmethod
    def _infer_document_type(*, title: str, content: str) -> str:
        text = f"{title}\n{content}"
        if any(keyword in text for keyword in ("校规", "学生管理", "学生手册", "考勤", "处分", "旷课")):
            return "school_rule"
        if "合同" in text:
            return "contract"
        if any(keyword in text for keyword in ("制度", "规定", "办法", "细则")):
            return "policy"
        if "手册" in text:
            return "manual"
        return "other"

    @classmethod
    def extract_uploaded_content(cls, *, filename: str, content_type: str | None, data: bytes) -> tuple[str, str]:
        normalized_name = filename.strip()
        lowered_name = normalized_name.lower()
        extension = "." + lowered_name.rsplit(".", 1)[-1] if "." in lowered_name else ""
        if extension not in cls.SUPPORTED_UPLOAD_EXTENSIONS and (content_type or "") not in cls.SUPPORTED_UPLOAD_MIME_TYPES:
            raise ValueError("仅支持 PDF 或 Word（.docx）文件")
        if len(data) == 0:
            raise ValueError("上传文件为空")
        if len(data) > cls.MAX_UPLOAD_BYTES:
            raise ValueError("文件过大，请上传 100MB 以内的 PDF 或 Word 文件")

        if extension == ".pdf" or content_type == "application/pdf":
            return normalized_name, cls._extract_pdf_text(data)
        return normalized_name, cls._extract_docx_text(data)

    @staticmethod
    def _extract_pdf_text(data: bytes) -> str:
        try:
            from pypdf import PdfReader
        except Exception as exc:  # pragma: no cover - depends on runtime package installation
            raise RuntimeError("当前环境缺少 pypdf，无法解析 PDF") from exc

        try:
            reader = PdfReader(BytesIO(data))
            parts = [(page.extract_text() or "").strip() for page in reader.pages]
        except Exception as exc:
            raise ValueError("PDF 解析失败，请确认文件未加密且内容可复制") from exc
        text = "\n\n".join(part for part in parts if part)
        if not text.strip():
            raise ValueError("PDF 中未提取到可检索文本")
        return text

    @staticmethod
    def _extract_docx_text(data: bytes) -> str:
        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover - depends on runtime package installation
            raise RuntimeError("当前环境缺少 python-docx，无法解析 Word 文件") from exc

        try:
            document = Document(BytesIO(data))
        except (BadZipFile, Exception) as exc:
            raise ValueError("Word 文件解析失败，请确认上传的是 .docx 文件") from exc
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        text = "\n".join([*paragraphs, *table_cells])
        if not text.strip():
            raise ValueError("Word 文件中未提取到可检索文本")
        return text

    async def list_documents(
        self,
        db: AsyncSession,
        *,
        agent_id: uuid.UUID,
        team_id: uuid.UUID,
    ) -> List[KnowledgeDocumentRead]:
        chunk_counts_subquery = (
            select(KnowledgeChunk.document_id, func.count(KnowledgeChunk.id).label("chunk_count"))
            .where(KnowledgeChunk.agent_id == agent_id, KnowledgeChunk.team_id == team_id)
            .group_by(KnowledgeChunk.document_id)
            .subquery()
        )
        stmt = (
            select(KnowledgeDocument, func.coalesce(chunk_counts_subquery.c.chunk_count, 0))
            .outerjoin(chunk_counts_subquery, KnowledgeDocument.id == chunk_counts_subquery.c.document_id)
            .where(KnowledgeDocument.agent_id == agent_id, KnowledgeDocument.team_id == team_id)
            .order_by(KnowledgeDocument.created_at.desc())
        )
        rows = await db.execute(stmt)
        return [self._read_model(document, chunk_count=int(chunk_count or 0)) for document, chunk_count in rows.all()]

    async def delete_document(
        self,
        db: AsyncSession,
        *,
        agent_id: uuid.UUID,
        team_id: uuid.UUID,
        document_id: uuid.UUID,
    ) -> bool:
        doc_stmt = select(KnowledgeDocument).where(
            KnowledgeDocument.id == document_id,
            KnowledgeDocument.agent_id == agent_id,
            KnowledgeDocument.team_id == team_id,
        )
        document = (await db.execute(doc_stmt)).scalar_one_or_none()
        if document is None:
            return False
        await db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == document_id))
        await db.delete(document)
        await db.commit()
        return True

    async def search(
        self,
        db: AsyncSession,
        *,
        agent_id: uuid.UUID,
        team_id: uuid.UUID,
        query: str,
        limit: int = 5,
        retrieval_mode: str = "optional_hybrid",
        article_no: str | None = None,
        include_near_misses: bool = True,
        document_type: str | None = None,
    ) -> List[KnowledgeSearchResult]:
        normalized_query = query.strip().lower()
        query_tokens = self._filter_query_tokens(self._tokenize(query))
        requested_article_no = article_no or self.extract_article_no(query)
        if not query_tokens and not requested_article_no:
            return []

        conditions = [KnowledgeChunk.agent_id == agent_id, KnowledgeChunk.team_id == team_id]
        if document_type:
            conditions.append(KnowledgeDocument.document_type == document_type)
        stmt = (
            select(KnowledgeChunk, KnowledgeDocument.document_type)
            .join(KnowledgeDocument, KnowledgeDocument.id == KnowledgeChunk.document_id)
            .where(*conditions)
        )
        rows = await db.execute(stmt)
        chunk_rows = rows.all()
        chunks = [chunk for chunk, _document_type in chunk_rows]
        document_types = {chunk.id: document_type or "other" for chunk, document_type in chunk_rows}

        return [
            KnowledgeSearchResult(
                document_id=chunk.document_id,
                chunk_id=chunk.id,
                title=chunk.title,
                content=chunk.content,
                score=ranked.score,
                match_type=ranked.match_type,
                document_type=document_types.get(chunk.id, "other"),
                article_no=chunk.article_no or self.extract_article_no(chunk.content),
                article_label=chunk.article_label or self.extract_article_label(chunk.content),
                section_path=chunk.section_path or [chunk.title, chunk.article_label or self.extract_article_label(chunk.content) or f"片段 {chunk.chunk_index + 1}"],
                page_no=chunk.page_no,
                citation_label=f"{chunk.title} / {chunk.article_label or self.extract_article_label(chunk.content) or f'片段 {chunk.chunk_index + 1}'}",
                is_direct_evidence=ranked.is_direct_evidence,
            )
            for ranked in knowledge_retrieval_ranker.rank(
                chunks=chunks,
                query=normalized_query,
                query_tokens=query_tokens,
                retrieval_mode=retrieval_mode,
                requested_article_no=requested_article_no,
                include_near_misses=include_near_misses,
                extract_article_no=self.extract_article_no,
                token_weight=self._token_weight,
            )[:limit]
            for chunk in [ranked.chunk]
        ]

    async def build_context(
        self,
        db: AsyncSession,
        *,
        agent_id: uuid.UUID,
        team_id: uuid.UUID,
        query: str,
        limit: int = 4,
    ) -> str:
        results = await self.search(db, agent_id=agent_id, team_id=team_id, query=query, limit=limit)
        if not results:
            return ""
        blocks = [
            f"[{index}] {item.title}\n{item.content}"
            for index, item in enumerate(results, start=1)
        ]
        return "\n\n".join(blocks)

    @staticmethod
    def _read_model(document: KnowledgeDocument, *, chunk_count: int) -> KnowledgeDocumentRead:
        return KnowledgeDocumentRead(
            id=document.id,
            agent_id=document.agent_id,
            title=document.title,
            content=document.content,
            chunk_count=chunk_count,
            document_type=getattr(document, "document_type", "other") or "other",
            source_filename=getattr(document, "source_filename", None),
            source_mime_type=getattr(document, "source_mime_type", None),
            source_hash=getattr(document, "source_hash", None),
            version_label=getattr(document, "version_label", None),
            effective_from=KnowledgeService._as_iso(document.effective_from) if getattr(document, "effective_from", None) else None,
            effective_to=KnowledgeService._as_iso(document.effective_to) if getattr(document, "effective_to", None) else None,
            status=getattr(document, "status", "ACTIVE") or "ACTIVE",
            metadata=getattr(document, "document_metadata", None) or {},
            created_at=KnowledgeService._as_iso(document.created_at),
            updated_at=KnowledgeService._as_iso(document.updated_at),
        )


knowledge_service = KnowledgeService()
